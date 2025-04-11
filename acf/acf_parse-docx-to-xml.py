from docx import Document
import sys
import os
import re
from lxml import etree
from fuzzywuzzy import process
from datetime import datetime
import pickle
import json
from collections import defaultdict
import urllib.parse
import requests
import csv
import pandas as pd

def find_source_link(cell, target_text, preceding_target_text=None):
    # Get the XML for the cell
    cell_xml = cell._element

    # Prep target text for searching
    orig_text = target_text
    target_text = target_text.strip().replace(' ', '').replace(')', '').replace('(', '').lower()

    # Prep preceding target text if provided
    if preceding_target_text is not None:
        preceding_target_text = preceding_target_text.strip().replace(' ', '').replace(')', '').replace('(', '').lower()

    preceding_text_found = False

    # Find all <w:p> elements in the cell
    paragraphs = cell_xml.findall('.//w:p', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

    # Iterate through <w:p> elements
    for paragraph in paragraphs:
        # Extract all <w:t> (text) elements in the paragraph
        text_elements = paragraph.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

        # Concatenate text from all <w:t> elements and clean it
        paragraph_text = ''.join([t.text.strip().replace(' ', '').replace('(', '').replace(')', '').replace('\u200b', '').replace("\u2009", "").replace("’", "'").lower() for t in text_elements if t.text])

        # Check for preceding target text if needed
        if preceding_target_text and not preceding_text_found:
            if preceding_target_text == paragraph_text:
                preceding_text_found = True

        # If preceding text was found (or not required), check for target text
        if preceding_text_found or preceding_target_text is None:
            # Find all <w:hyperlink> elements in the paragraph
            hyperlinks = paragraph.findall('.//w:hyperlink', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

            for hyperlink in hyperlinks:
                # Extract all <w:t> elements in the hyperlink
                link_text_elements = hyperlink.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                full_text = ''.join([t.text.strip().replace(' ', '').replace('\u200b', '').replace("\u2009", "").replace("’", "'").replace('(', '').replace(')', '').lower() for t in link_text_elements if t.text])

                # Match the full hyperlink text with target_text
                if target_text == full_text:
                    # Extract rID from hyperlink
                    r_id = hyperlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                    anchor = hyperlink.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}anchor")

                    if r_id:
                        rel = cell.part.rels.get(r_id)
                        if rel:
                            # Build the full URL
                            full_target = rel._target
                            if anchor:
                                full_target += f"#{anchor}"  # Append anchor if present

                            return full_target

    # OPTION 2: LINKS ARE STORED IN <w:instrText> TAGS, NESTED INSIDE <p> tags alongside text

    # Loop through <w:p> tags; check if the paragraph contains the target text in any <w:t> tag
    for para in cell.paragraphs:
        paragraph_text = "".join(run.text.strip().lower() for run in para.runs).replace(' ', '')

        # If we find the paragraph containing our target text, look for <w:instrText>
        if target_text in paragraph_text.replace('’', "'"):
            for run in para.runs:
                for element in run._element:
                    # Check if the element is <w:instrText>
                    if element.tag.endswith('instrText') and 'HYPERLINK' in element.text:
                        # Extract the URL between the first pair of quotes
                        parts = element.text.split('"')
                        if len(parts) > 1:
                            hyperlink_url = parts[1]
                            
                            # Check for the presence of the \l switch and append as an anchor
                            if "\\l" in element.text:
                                # Find the text following \l
                                l_index = element.text.index("\\l") + 2
                                anchor_text = element.text[l_index:].strip()
                                if anchor_text.startswith('"') and anchor_text.endswith('"'):
                                    anchor_text = anchor_text[1:-1]  # Remove surrounding quotes
                                hyperlink_url += f"#{anchor_text}"

                        return hyperlink_url    

    # If no matching hyperlink was found, return empty string (no None so that we avoid string processing issues downstream)
    print(f"\n\nWARNING: Word doc does not appear to include a hyperlink for this string: {orig_text}")
    try:
        print(f'\tTRIED: {full_text}')
    except UnboundLocalError:
        pass
    try:
        print(f'\tTRIED: {paragraph_text}')
    except UnboundLocalError:
        pass
    
    return None

def find_closest_match(user_term, vocabulary, threshold=90):

    match, score = process.extractOne(user_term, vocabulary)
    return match if score >= threshold else None

def write_error(details, error_msg):
    # set up the audit log where we will record error

    with open(details['tmp_audit_log'], 'a', encoding='utf-8') as fo:
        fo.write(f"{error_msg}\n") 

def prep_cell_text(raw_txt):

    #clean string--remove curly quotes and zero-width spaces
    txt_step_1 = raw_txt.replace('’', "'").replace('“', '"').replace('”', '"').replace('\u200b', '').replace("\u2009", " ").replace('•', '').replace('⦁', '')

    #split into list of strings, trim white space from each string
    txt_step_2 = [item.strip() for item in txt_step_1.splitlines()]

    # had some instances of ':' sneaking into data; could be other issues. Let's catch and remove those
    txt_step_3 = [item for item in txt_step_2 if item not in [':']]

    #remove any empty/blank strings
    final_txt = [item for item in txt_step_3 if item.strip()]

    return final_txt

def check_controlled_vocabs(cv_list, cv_used, details, current_position):

    # controlled vocabs
    cvs = {
        "offices": [
            "Administration on Children, Youth, and Families (ACYF)",
            "Children's Bureau (CB)",
            "Family and Youth Services Bureau (FYSB)",
            "Office of Child Support Enforcement (OCSE)",
            "Office of Family Assistance (OFA)",
            "Office of Head Start (OHS)",
            "Office of Planning, Research, and Evaluation (OPRE)",
            "Office of Child Care (OCC)",
            "Office of Community Services (OCS)",
            "Office of Early Childhood Development (ECD)",
            "Office on Trafficking in Persons (OTIP)",
            "Administration for Native Americans (ANA)",
            "Office of Refugee Resettlement (ORR)",
            "N/A"
        ],
        "domains": [
            "Public Records",
            "Medical Assistance",
            "Public Assistance",
            "Child Welfare Services"
        ],
        "federal": ["21st Century Cures Act", "Adam Walsh", "BJS", "CAPEA", "CAPTA", "CCDBG", "CSBG", "EETC", "Evidence Act/CIPSEA", "FCIA", "FERPA", "FISMA", "FVPSA", "Head Start Act", "HIPAA/HITECH", "ICWA", "IDEA", "IRS", "LIHEAP", "Medicare Act/CMS", "MVHAA", "NCHS", "NAPA", "OASDI", "Privacy Act", "Refugee Education Assistance Act", "RHYA", "SAMHSA/SAPT", "Section 1137 of the SSA", "SNAP/2018 Farm Bill", "SSA Title IV-A", "SSA Title IV-B", "SSA Title IV-D", "SSA Title IV-E", "SSA Title XIX", "SSBG", "TVPA", "UIFSA", "US Repatriation Program", "N/A"],
        "terms": ["Abuse and Neglect", "Adoption and Foster Care", "Background Checks", "Biometric Information", "Breach Response", "Child Care", "Child Support", "Children and Youth Services", "Confidentiality", "Criminal Justice/Courts", "Cross-jurisdictional", "Data Collection", "Data Recipient Requirements", "Data Retention", "Databases", "Disability", "Domestic Violence", "Early Childhood", "Economic Security and Mobility", "Education – Higher Ed", "Education – K-12", "Education – Pre-K", "Family Services", "Genetic Information", "Grants & Funding", "Health Care", "How the Law Relates to Other Laws", "Human Trafficking", "Immigration & Refugee Services", "Income Verification", "Individual Subject Rights", "Information Security", "Information Technology Systems", "Medicaid/Medicare", "Mental Health", "Military", "Minorities", "Missing and Unidentified Persons", "Non-Compliance and Consequences for Misuse", "Nutrition Assistance", "Parents", "Preservation of Culture", "Program Eligibility", "Runaway and Homeless Youth", "Substance Abuse", "Taxpayer Information", "Tribal/Native American", "Use and Sharing – Programmatic", "Use and Sharing – Research"]
    }
    
    # first, make sure no duplicates are included
    no_duplicates = []
    duplicates = []

    for term in cv_list:
        if term in no_duplicates:
            duplicates.append(term)
        else:
            no_duplicates.append(term)

    # log duplicate entry error(s) and use 'no duplicate list'
    if duplicates:
        for dup in duplicates:
            write_error(details, f"{current_position} - Duplicate values: '{dup}' (from {cv_used.capitalize()} vocabulary) included multiple times.")
        cv_list = no_duplicates
    
    # now review instances where a provided term is not found in our vocabs
    # first, ID the index of any term that is not included
    problems = [ term for term in cv_list if term not in cvs[cv_used] ]

    # if we do have any problematic terms, we'll use fuzzywuzzy (via our find_closest_match function) to identify the closest match, if possible
    if problems:
        for term in problems:
            #if any empty strings made their way in, get rid of them
            if not term:
                cv_list.remove(term)
                continue 

            # get the term's index
            idx = cv_list.index(term)
            
            # use a regex to find the closest match to our term
            closest_match = find_closest_match(term, cvs[cv_used])
            
            # if a term is proposed, log our correction and then replace the problematic one with the proposed term
            if closest_match is not None:
                write_error(details, f"{current_position} - Incorrect term: MS Word included '{term}'; replaced with '{closest_match}' (from {cv_used.capitalize()} vocabulary).")
                cv_list[idx] = closest_match
                continue

            # If we have no match, do not include term in XML; log error instead.
            else:
                write_error(details, f"{current_position} - Unidentified term: MS Word included '{term}', which does not exist in the {cv_used.capitalize()} vocabulary). Revise Word doc and re-run XML generation, if needed.")
                cv_list.pop(idx)
                print(f"\n\nWARNING: {term} will not be included in XML, as it is not part of the {cv_used.capitalize()} vocabulary. See {current_position} to determine if .docx should be corrected.")

    # NOTE: for the 'term' vocab, additional terms may need to be applied; add terms if needed and then log error
    if 'terms' in cv_used:
        if ('Child Support' in cv_list) and ('Economic Security and Mobility' not in cv_list):
            cv_list.append('Economic Security and Mobility')
            write_error(details, f"{current_position} - Missing term: when 'Child Support' is used, 'Economic Security and Mobility' must also be included.")

    return cv_list

def parse_requirement_blocks(data, cell, temp_list, details, parent_position):

    # Patterns to identify the different parts of each entry
    label_desc_state_code_pattern = rf"{details['state_code_pattern']}"
    entities_pattern = re.compile(r"Who Law Applies To:(.*)")
    tags_pattern = re.compile(r"(?:Tags|Tag\(s\)|Tags\(s\)):(.*)")
    
    #set index counter
    i = 0

    data_limit = len(data)

    # we will loop through data in slices of 3; continue looping until we've reached the end of the data
    while i < data_limit:
    
        # Process three items starting from the current index
        data_slice = data[i:i+3]

        record = {}

        match = re.search(label_desc_state_code_pattern, data_slice[0])
        if match:
            label = match.group(1).strip()
            description = match.group(2).strip()
            state_code = match.group(3).strip()

        else:
            try:
                dash = match_dash(data_slice[0])
                label, description = data_slice[0].split(dash, 1) 
                label = label.strip()

                description, state_code = description.rsplit(' (', 1)
                description = description.strip()
                
                state_code = state_code.replace(')', '').strip()
            except ValueError:
                print(f'\nValueError: unable to parse this content: {data_slice[0]}') 

        try:
            # update dict
            record['label'] = label
            record['description'] = description
            record['state_code'] = state_code
            record['source'] = find_source_link(cell, state_code)

            if record['source'] is None:
                print('\tREQ LABEL', label)
                print("\tREQ CODE", state_code)

        except UnboundLocalError:
            print(f'\nUnboundLocalError: unable to parse this content: {data_slice[0]}')

        #update current position
        try:
            current_position = f"{parent_position} - {state_code}"
        except UnboundLocalError:
            print(data_slice)
            sys.exit(1)

        # Parse entities (as a list)
        try:
            entities_match = entities_pattern.match(data_slice[1])
            if entities_match:
                entities_text = entities_match.group(1)
                record['entities'] = [entity.strip() for entity in entities_text.split(";")]
        except IndexError:
            print(data_slice)
            sys.exit(1)
        
        # Parse tags (as a list)
        tags_match = tags_pattern.match(data_slice[2])
        if tags_match:
            tags_text = tags_match.group(1)
        else:
            #print(data_slice)
            tags_text = data_slice[2].split(':', 1)[1]

        # check our terms and assign to dict
        all_tags = [tag.strip() for tag in re.split(r"[;,:]", tags_text)]
        record['tags'] = check_controlled_vocabs(all_tags, "terms", details, current_position)

        temp_list.append(record)

        # increment to move on to the next group
        if data_limit - i < 3:
            i += (data_limit - i)
        else:
            i += 3

    return temp_list

def parse_to_dict(list_of_strings, cell, search_term):

    # we need to be able to handle lists of search_terms
    if not isinstance(search_term, list):
        search_term = [ search_term ]
    
    for t_idx, t in enumerate(list_of_strings):

        #the search terms should always occur at beginning of string
        if any(t.lower().startswith(term.lower()) for term in search_term):

            # if found, assign values to dictionary
            temp_dict = {
                "current_position": t.strip(),
                "name": list_of_strings[t_idx+1],
                "source": find_source_link(cell, list_of_strings[t_idx+1], t),
                "start_idx": t_idx,
                "end_idx": t_idx+1
            }

            # NOTE: we have at least one state that does not include Title #'s. Add handling for this...
            try:
                temp_dict['number'] = t.split(' ', 1)[1].strip()
            except IndexError:
                temp_dict['number'] = None

            # we also need to get a form of the name to use as a dict key. Just do this for everything, but it's only used with Titles. We have to take additional steps in case this state doesn't include Title Numbers
            if temp_dict['current_position'].lower() == search_term[0].lower():
                temp_dict['current_position'] = f"{temp_dict['current_position']} {temp_dict['name']}"
                temp_dict['title_key'] = temp_dict["current_position"]
            else:
                temp_dict['title_key'] = f"{temp_dict["current_position"]}-{temp_dict["name"]}".replace(' ', '').lower()

            #note that if this record uses multiple 'Part' names, we have to see if the altName (i.e., the second partName) is used
            if len(search_term) > 1:
                if search_term[1] in temp_dict['current_position']:
                    temp_dict['altName'] = search_term[1]

            # exit for loop as soon as we found our match
            break

    try:
        return temp_dict
    except NameError:
        return None

def return_index(row_cells, target_text):

    longest_string_idx = 0
    longest_string_len = 0
    
    for i in range(len(row_cells)):
        if target_text in row_cells[i].text:
            if len(row_cells[i].text) > longest_string_len:
                longest_string_len = len(row_cells[i].text)
                longest_string_idx = i 

    if longest_string_len > 0:
        return longest_string_idx
    else:
        return None

def match_dash(text_string):

    for punc in [' – ', '–', ' - ', '-', '—']:
        if punc in text_string:
            return punc

    return None

def parse_tables(doc, details, record_data, object_type):

    # Loop through each table in the document
    for i, table in enumerate(doc.tables):

        # skip any table with only 1 column in first row
        if (len(table.rows[0].cells) == 1) or ("Table of Contents" in table.rows[0].cells[0].text) or ("Domain Color Coding" in table.rows[0].cells[0].text):
            continue

        # Loop through each row in the table
        for row_idx, row in enumerate(table.rows):

            # initial process for getting Title info
            if object_type == "title":

                # Check if the first cell of the row is empty; a title row should always be empty
                initial_cell_text = row.cells[0].text.strip()
                if not initial_cell_text:

                    #some rows will actually have more than 3 cells: need to verify where indexes will start
                    if len(row.cells) > 3:
                        title_idx = return_index(row.cells, details['titleName'])
                        office_idx = return_index(row.cells, 'ACF Offices Associated')              
                    else:
                        title_idx = 1
                        office_idx = 2 

                    # get the text in the adjacent cell; split lines and loop through text, checking for regex match. NOTE: simple string matching might also work...
                    cell_text = prep_cell_text(row.cells[title_idx].text)
                    title_dict = parse_to_dict(cell_text, row.cells[title_idx], details['titleName'])

                    # Assign our dict values
                    if title_dict is not None:
                        current_title_key = title_dict['title_key']
                        record_data[current_title_key] = {
                            "name": title_dict['name'],
                            "number": title_dict['number'],
                            "source": title_dict['source']
                        }
                            
                        #set 'current position' so we can track any errors
                        current_position = title_dict['current_position']
                    
                        #if record uses 'categories', get info:
                        if details["category"]:
                            # category always occurs before the title; if the title is at index 0, something is fishy...
                            if title_dict['start_idx'] == 0:
                                print('\n\nWARNING: config indicates that "Categories" are used, but none found in title cell!!')
                            else:
                                category_name = cell_text[title_dict['start_idx']-1].strip()
                                if 'Part' in category_name:
                                    source_text = category_name.split(': ', 1)[1]
                                else:
                                    source_text = category_name
                                
                                category_source = find_source_link(row.cells[title_idx], source_text)
                            
                                #update 'current position'
                                current_position = f"{category_name} - {current_position}"

                                # Assign our dict values
                                record_data[current_title_key]['category'] = {
                                    "name": category_name
                                }

                                if category_source is not None:
                                    record_data[current_title_key]['category']['source'] = category_source

                        # 'ACF Offices Associated' should be in the cell immediately next to the titleName; get offices and add to dict
                        if office_idx is not None:
                            office_txt = prep_cell_text(row.cells[office_idx].text.replace('ACF Offices Associated', ''))
                            
                            # check term list for errors, then add to our dictionary
                            record_data[current_title_key]['officesAssociated'] = check_controlled_vocabs(office_txt, "offices", details, current_position)

            # additional process to add Article info
            elif object_type == "article":

                # all rows with Article information should have Domain information at the very top of the first cell
                if 'Domain' in row.cells[0].text:

                    # set variables
                    overview_cell = row.cells[0]

                    #be ready for 'titleContent' scenario
                    found_titleContent = False

                    # prepare our text
                    article_overview = prep_cell_text(overview_cell.text)

                    # determine which title we are working with
                    title_dict = parse_to_dict(article_overview, overview_cell, details['titleName'])

                    if title_dict is not None:
                        current_title = title_dict['current_position']
                        current_title_key = title_dict['title_key']
                        title_end_idx = title_dict['end_idx']

                    #set current position; will need to see if we have a category
                    try:
                        if record_data[current_title_key].get('category'):
                            current_position = f"{record_data[current_title_key]['category']['name']} - {current_title}"
                        else:
                            current_position = current_title
                    except: 
                        print('\n\nWARNING: may have a problem with "current_title_key" variable--check for any variation in Title name:\n\n', article_overview)
                        sys.exit(1)

                    # add 'articles' list to our dict if it's not already there
                    record_data[current_title_key].setdefault('articles', [])

                    # get article info
                    if details['articleName']:
                        temp_article_dict = parse_to_dict(article_overview, overview_cell, details['articleName'])
                    else:
                        temp_article_dict = None

                    # Provide warning if we didn't get an article; continue on to next item
                    if temp_article_dict is None:
                        if not details.get('titleContent', False):
                            print(f"\n\nWARNING: failed to ID article; review cell contents:\n\n\t{'\n\t'.join(article_overview)}")
                            continue
                        # if we anticipate 'titleContent', assume we have found it!
                        else: 
                            found_titleContent = True
                            temp_article_dict = title_dict
                            temp_article_dict['found_titleContent'] = True

                    # now that we have our temp_article_dict, get the domain and check for any errors; the domain name should always be at index [1]
                    article_domain = check_controlled_vocabs([article_overview[1]], "domains", details, current_position)
                    if not article_domain:
                        print(f'\n\nWARNING: incorrect domain; review cell contents: {article_overview[0:2]}')
                    else:
                        temp_article_dict['domain'] = article_domain[0]

                    # if we expect to find subtitles, check to see if there is any text between the Title and the Article
                    if details['subtitleName']:
                        subtitle_slice = article_overview[title_end_idx+1:temp_article_dict['start_idx']]

                        # if there is actually text here, return subtitle info
                        if len(subtitle_slice) > 0:
                            temp_article_dict['subtitle'] = parse_to_dict(subtitle_slice, overview_cell, details['subtitleName'])

                            #now set current position
                            if temp_article_dict['subtitle']['current_position'] not in current_position:
                                current_position += f" - {temp_article_dict['subtitle']['current_position']} - {temp_article_dict['current_position']}"
                    else:
                        if temp_article_dict['current_position'] not in current_position:
                            current_position += f" - {temp_article_dict['current_position']}"

                    # get associated federal records; then check for errors
                    try:  
                       index_of_associated_records = [i for i, txt in enumerate(article_overview) if 'Associated Federal Records' in txt][0]
                    except IndexError:
                        print(f'\n\nWARNING: "Associated Federal Records" is missing at {current_position}. Check Word DOCX and retry.')
                        sys.exit(1)

                    temp_article_dict['associatedFederalRecords'] = check_controlled_vocabs(article_overview[index_of_associated_records + 1:], "federal", details, current_position)

                    # if we expect to find parts, check to see if there is any text between the Article and the Associated Federal Records
                    if details['partName']:
                        part_slice = article_overview[temp_article_dict["end_idx"]+1:index_of_associated_records]

                        # if there is actually text here, return part info
                        if len(part_slice) > 0:
                            temp_article_dict['part'] = parse_to_dict(part_slice, overview_cell, details['partName'])

                            #NOTE: we sometimes have issues with the parsing; exit if we have an error so that we can figure out the issue
                            try:
                                if temp_article_dict['part']['current_position'] not in current_position:
                                    current_position += f" - {temp_article_dict['part']['current_position']}"
                            except TypeError:
                                print(part_slice)
                                sys.exit(1)

                            # we will only have a sub-part if there is a part
                            if details['subPartName']:

                                part_end_idx = article_overview.index(temp_article_dict['part']['name']) + 1

                                #the subPart will occur between the Part and Associated Records
                                subPart_slice = article_overview[part_end_idx:index_of_associated_records]

                                if len(subPart_slice) > 0:
                                    temp_article_dict['part']['subPart'] = parse_to_dict(subPart_slice, overview_cell, details['subPartName'])

                                    if temp_article_dict['part']['subPart']['current_position'] not in current_position:
                                        current_position += f" - {temp_article_dict['part']['subPart']['current_position']}"

                    # now move to the next column and get definitions and requirements. Note that some Word Docx files may have a varied # of cells per row. We need to test this... 
                    statute_cell_index = 1
                    while True:
                        if prep_cell_text(row.cells[0].text) != prep_cell_text(row.cells[statute_cell_index].text):
                            break
                        else:
                            statute_cell_index += 1

                    statutes_cell = row.cells[statute_cell_index]
                    statutes = prep_cell_text(statutes_cell.text)
                            
                    #get indices for Definitions and Requirements
                    def_index = None
                    req_index = None

                    # set up lists to carry definition and requirement information
                    definitions = []
                    requirements = []

                    for index, item in enumerate(statutes):
                        if any(item.lower().startswith(phrase) for phrase in ['definitions related to', 'definitions for ']):
                            def_index = index
                            break

                    for index, item in enumerate(statutes):
                        if any(item.lower().startswith(phrase) for phrase in ['requirements related ', 'requirements for ', 'requirements regarding ', 'regulations regarding ']):
                            req_index = index
                            # NOTE: we need to edit Word doc to replace 'Regulations' with 'Requirements'; give a warning if so
                            if "Regulations" in statutes[index]:
                                print(f'\n\nWARNING: Found REGULATION at {current_position}')
                            break

                    # we may not have definitions; make sure we found them
                    if def_index is not None:
                        #establish the index to slice our list of strings
                        start_index = def_index+1

                        if req_index is None:
                            def_info = statutes[start_index:]
                        else:
                            def_info = statutes[start_index:req_index]

                        defn_pattern = rf"{details['statute_pattern']}"

                        # loop through any definitions
                        for defn in def_info:
                            
                            defn_match = re.search(defn_pattern, defn)
                            if defn_match:
                                state_code = defn_match.group(1)
                                terms = defn_match.group(2)
                            else:
                                dash = match_dash(defn)

                                try:
                                    state_code, terms = defn.split(dash, 1)
                                except UnboundLocalError:
                                    print(f'\n\nWARNING - unable to split this text: {defn}')
                                    sys.exit(1)

                            # add state code, source, and defined terms to a dictionary; append to our definitions
                            temp_defn_dict = { 
                                "state_code": state_code.strip(),
                                "source": find_source_link(statutes_cell, state_code),
                                "defined_terms": [t.strip() for t in terms.split(',')]
                            }

                            if temp_defn_dict['source'] is None:
                                print('\tDEFINITION:', defn)
                            definitions.append(temp_defn_dict)


                    # add definitions info to article dict; if none have been found, we just have an empty list
                    temp_article_dict['definitions'] = definitions

                    # Now, verify that we have requirements; if so, use index to pull out all requirement text
                    if req_index is not None:

                        # There may be an empty line between the 'Requirements' statement and the statutes. This code makes sure we are not including a blank line at the start
                        check_index = req_index+1
                        while True:
                            if len(statutes[check_index]) > 0:
                                break
                            else:
                                check_index += 1

                        # slice our original list to include all the statute information
                        req_info = statutes[check_index:]

                        # parse out requirements and add to temp dict
                        requirements = parse_requirement_blocks(req_info, statutes_cell, requirements, details, current_position)
                        temp_article_dict['requirements'] = requirements

                    if temp_article_dict not in record_data[current_title_key]['articles']:
                        record_data[current_title_key]['articles'].append(temp_article_dict)

    return record_data[current_title_key]

def validate_xml(xml_file, xsd_file):
    # Parse the XML file
    with open(xml_file, 'r') as xml:
        xml_doc = etree.parse(xml)

    # Parse the XSD file
    with open(xsd_file, 'r') as xsd:
        xsd_doc = etree.parse(xsd)

    # Create an XMLSchema object
    xml_schema = etree.XMLSchema(xsd_doc)

    # Validate the XML file against the schema
    is_valid = xml_schema.validate(xml_doc)

    if not is_valid:
        print("\n\nXML is not valid. Errors:")
        # read in xml to be able to print specific lines
        with open(xml_file, 'r', encoding='utf-8') as file:
            xml_lines = file.readlines()

        #collect all our errors
        xml_errors = []
        for error in xml_schema.error_log:
            print(f'\n\n - Line:', error.line)
            print(f' - Message:', error.message)

def write_xml(details, record_data):
    # Declare the namespace URI
    namespace = {'xsi': 'http://www.w3.org/2001/XMLSchema-instance'}
    
    # Create the root element <records>
    root = etree.Element("record", nsmap=namespace)

    # Add the xsi:noNamespaceSchemaLocation attribute
    root.set("{http://www.w3.org/2001/XMLSchema-instance}noNamespaceSchemaLocation", "../schema_final.xsd")

    #set up top-level elements
    etree.SubElement(root, "state").text = details.get('state', '')
    etree.SubElement(root, "articleName").text = details.get('articleName', '')
    etree.SubElement(root, "titleName").text = details.get('titleName', '')
    if details.get('partName', []):
        etree.SubElement(root, "partName").text = details['partName'][0]
    if details.get('subPartName'):
        etree.SubElement(root, "subPartName").text = details.get('subPartName', '')
    if details.get('subtitleName'):
        etree.SubElement(root, "subtitleName").text = details.get('subtitleName', '')

    # set a dummy value to check current category
    current_category = 'none'

    for title in record_data.keys():
        # include 'category' info, if applicable
        if record_data[title].get('category', {}):
            # only add 'category' element if it's a new one. Compare value to 'current category'
            if record_data[title]['category'].get('name', '') not in current_category:
                cat_elem = etree.SubElement(root, "category")
                etree.SubElement(cat_elem, "name").text = record_data[title]['category'].get('name', '')
                if record_data[title]['category'].get('source'):
                    etree.SubElement(cat_elem, "source").text = record_data[title]['category'].get('source', '')

                #update 'current category'
                current_category = record_data[title]['category']['name']
                
            # NOTE: if we have a category, then Title will be nested under it. Otherwise, Title will be a child of root
            title_elem = etree.SubElement(cat_elem, "title")
        else:
            title_elem = etree.SubElement(root, "title")
        
        # Add main title elements
        etree.SubElement(title_elem, "number").text = record_data[title].get('number', '')
        etree.SubElement(title_elem, "name").text = record_data[title].get('name', '') 
        etree.SubElement(title_elem, "source").text = record_data[title].get('source', '')

        # Add officesAssociated elements
        offices_elem = etree.SubElement(title_elem, "officesAssociated")
        for office in record_data[title].get("officesAssociated", []):
            etree.SubElement(offices_elem, "office").text = office

        # Loop through articles within each title
        for article in record_data[title].get("articles", []):

            # check if this is a case with 'titleContent'; default is False
            if article.get("found_titleContent", False):
                article_elem = etree.SubElement(title_elem, "titleContent")
                etree.SubElement(article_elem, "domain").text = article.get("domain", '')
            
            # if not, set up article and write domain
            else:
                article_elem = etree.SubElement(title_elem, "article")
                etree.SubElement(article_elem, "domain").text = article.get("domain", '')

            # add subtitle info, id present
            if article.get("subtitle", {}):
                subtitle_elem = etree.SubElement(article_elem, "subtitle")
                etree.SubElement(subtitle_elem, "number").text = article['subtitle'].get("number", "")
                etree.SubElement(subtitle_elem, "name").text = article['subtitle'].get("name", "")
                etree.SubElement(subtitle_elem, "source").text = article['subtitle'].get("source", "")

            # Write basic article info--only if this is NOT a 'titleContent scenario'
            if not article.get("found_titleContent", False):
                etree.SubElement(article_elem, "number").text = article.get("number", '')
                etree.SubElement(article_elem, "name").text = article.get("name", '')
                etree.SubElement(article_elem, "source").text = article.get("source", '')

            # Add part info, if applicable
            if article.get('part', {}):
                part_elem = etree.SubElement(article_elem, "part")
                etree.SubElement(part_elem, "number").text = article['part'].get("number", '')
                etree.SubElement(part_elem, "name").text = article['part'].get("name", '')
                etree.SubElement(part_elem, "source").text = article['part'].get("source", '')
                
                # determine if the part used an 'altName'
                if article['part'].get('altName'):
                    etree.SubElement(part_elem, "altName").text = article['part'].get("altName", '')

                # check to see if there is a subPart
                if article['part'].get('subPart', {}):
                    subPart_elem = etree.SubElement(part_elem, "subPart")
                    etree.SubElement(subPart_elem, "number").text = article['part']['subPart'].get("number", '')
                    etree.SubElement(subPart_elem, "name").text = article['part']['subPart'].get("name", '')
                    etree.SubElement(subPart_elem, "source").text = article['part']['subPart'].get("source", '')

            # Add associated federal records
            fedrecords_elem = etree.SubElement(article_elem, "associatedFederalRecords")
            for rec in article.get("associatedFederalRecords", []):
                etree.SubElement(fedrecords_elem, "federal").text = rec

            # Add definitions--only if we have at least one
            if article.get("definitions", []):
                def_elem = etree.SubElement(article_elem, "definitions")
                for defn in article.get("definitions"):
                    def_statute_elem = etree.SubElement(def_elem, "statute")
                    etree.SubElement(def_statute_elem, "stateCode").text = defn.get("state_code", '')
                    etree.SubElement(def_statute_elem, "source").text = defn.get("source", '')
                    
                    def_terms_elem = etree.SubElement(def_statute_elem, "definedTerms")
                    for def_term in defn.get("defined_terms", []):
                        etree.SubElement(def_terms_elem, "definedTerm").text = def_term

            # Add requirements
            if article.get("requirements", []):
                req_elem = etree.SubElement(article_elem, "requirements")
                for req in article.get("requirements", []):
                    req_statute_elem = etree.SubElement(req_elem, "statute")
                    etree.SubElement(req_statute_elem, "label").text = req.get("label", '')
                    etree.SubElement(req_statute_elem, "description").text = req.get("description", '')
                    etree.SubElement(req_statute_elem, "stateCode").text = req.get("state_code", '')
                    etree.SubElement(req_statute_elem, "source").text = req.get("source")

                    # Add appliesTo entities
                    applies_elem = etree.SubElement(req_statute_elem, "appliesTo")
                    for ent in req.get("entities", ['']):
                        etree.SubElement(applies_elem, "entity").text = ent

                    # Add terms tags
                    terms_elem = etree.SubElement(req_statute_elem, "terms")
                    for tag in req.get("tags", []):
                        #watch out for lxml double-escaping ampersands...
                        etree.SubElement(terms_elem, "term").text = tag.replace('&amp;', '&')

    # Create the ElementTree object
    tree = etree.ElementTree(root)

    # Write the XML to a file
    xml_file = os.path.join(details["out_dir"], f"{details['state'].lower().replace(' ', '_')}_{datetime.now().strftime("%Y%m%d")}.xml")

    with open(xml_file, "wb") as file:
        tree.write(file, pretty_print=True, xml_declaration=True, encoding="UTF-8")

    return xml_file

def generate_html(xml_file, xsl_file):

    # Load the XML file
    xml_tree = etree.parse(xml_file)

    # Load the XSL file
    xsl_tree = etree.parse(xsl_file)

    # Create an XSLT transformer
    transform = etree.XSLT(xsl_tree)

    # Apply the transformation to the XML data
    html_tree = transform(xml_tree)

    # Write the HTML to a file
    html_file = os.path.join(details["out_dir"], f"{details['state'].lower().replace(' ', '_')}_{datetime.now().strftime("%Y%m%d")}.html")
    html_tree.write(html_file, pretty_print=True, method="html")

def check_all_hyperlinks_from_docx(doc, xml_file):

    # First, get all the URLs in our XML
    tree = etree.parse(xml_file)
    root = tree.getroot()

    # Extract all hyperlinks in <source> tags at various levels of hierarchy
    xml_hyperlinks = set()
    
    # Find all <source> elements at different levels of the XML hierarchy
    for source in root.xpath('.//source'):
        if source.text:
            xml_hyperlinks.add(source.text.strip().lower())

    #Next, look through hyperlinks in MS docx and document which ones are not present
    # Dictionary to store hyperlinks; each URL maps to a set of anchor texts
    missing_docx_hyperlinks = {}

    # Iterate over every table in the document
    for t_i, table in enumerate(doc.tables):

         # skip any table with only 1 column in first row
        if (len(table.rows[0].cells) == 1) or ("Table of Contents" in table.rows[0].cells[0].text):
            continue

        for r_i, row in enumerate(table.rows):

            for c_i, cell in enumerate(row.cells):
                # Access the XML structure of the cell
                cell_xml = cell._element
                
                # Find all hyperlinks in the cell
                hyperlinks = cell_xml.findall('.//w:hyperlink', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                
                for hyperlink in hyperlinks:

                    # Get the rID and anchor from the hyperlink
                    r_id = hyperlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                    anchor = hyperlink.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}anchor")
                    
                    if r_id:
                        rel = cell.part.rels.get(r_id)
                        if rel:
                            # Build the full URL, including the anchor if present
                            full_target = rel.target_ref
                            if anchor:
                                full_target += f"#{anchor}"  # Append the anchor as a fragment

                            # if we missed this hyperlink in our XML, get more info
                            if full_target.lower() not in xml_hyperlinks:

                                # see if it's URL-encoded; if it is in our list, continue, otherwise, document error
                                if urllib.parse.unquote(full_target.lower()) in xml_hyperlinks:
                                    continue
                                if urllib.parse.quote(full_target.lower()) in xml_hyperlinks:
                                    continue

                                # Retrieve the text associated with the hyperlink
                                text_elements = hyperlink.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                                if len(text_elements) == 0:
                                    continue
                                
                                # Combine text elements (handling irregular text breaks)
                                full_text = ''.join([text.text for text in text_elements if text.text])

                                #now save info 
                                if not missing_docx_hyperlinks.get(full_target):
                                    missing_docx_hyperlinks[full_target] = []
                                missing_docx_hyperlinks[full_target].append(full_text)

                # Check for HYPERLINK fields in <w:p> elements
                paragraphs = cell_xml.findall('.//w:p', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                for paragraph in paragraphs:
                    # Find the <w:instrText> element
                    instr_text = paragraph.find('.//w:instrText', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

                    if instr_text is not None and 'HYPERLINK' in instr_text.text:
                        # Extract the hyperlink URL from the instrText
                        parts = instr_text.text.split('"')
                        if len(parts) > 1:
                            hyperlink_url = parts[1].lower()

                            if "\\l" in instr_text.text:
                                # Find the text following \l
                                l_index = instr_text.text.index("\\l") + 2
                                anchor_text = instr_text.text[l_index:].strip()
                                if anchor_text.startswith('"') and anchor_text.endswith('"'):
                                    anchor_text = anchor_text[1:-1]  # Remove surrounding quotes
                                hyperlink_url += f"#{anchor_text}".lower()

                            if hyperlink_url.lower() not in xml_hyperlinks:

                                # Extract all text in the same paragraph
                                text_elements = paragraph.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                                paragraph_text = ''.join([text.text.strip() for text in text_elements if text.text])

                                #now save info 
                                if not missing_docx_hyperlinks.get(hyperlink_url):
                                    missing_docx_hyperlinks[hyperlink_url] = []
                                missing_docx_hyperlinks[hyperlink_url].append(paragraph_text)

    # If there are missing hyperlinks, print them
    if missing_docx_hyperlinks:
        print(f"\n\nFound {len(missing_docx_hyperlinks)} hyperlinks in DOCX that are not in the XML:")
        for url, text in missing_docx_hyperlinks.items():
            print(f"\n\nURL: {url}") 
            for t in text:  
                print(f'\tTEXT: {t}')

def main(details):
    
    # load DOCX file
    doc = Document(details["input_doc"])

    # set array to capture info
    record_data = {}

    # get our title information
    print('\n\nGetting title information...')
    parse_tables(doc, details, record_data, "title")

    # get our article information
    print('\n\nGetting article information...')
    parse_tables(doc, details, record_data, "article")

    # now write to XML with lxml
    print('\n\nWriting XML...')
    xml_file = write_xml(details, record_data)

    # validate our XML
    print('\n\nValidating XML...')
    validate_xml(xml_file, details['xsd_file'])

    # verify that we retrieved all links from docx
    print('\n\nMaking sure all links are in XML...')
    check_all_hyperlinks_from_docx(doc, xml_file)

    # Finally, generate HTML; NOTE: in the future, add xsl_file path as variable to config
    print('\n\nGenerating HTML...')
    generate_html(xml_file, details['xsl_file'])

    #clean up our error log, if it exists; read in previous error messages and sort
    if os.path.exists(details['tmp_audit_log']):
        print('\n\nWriting audit log and cleaning up...')
        with open(details['tmp_audit_log'], 'r', encoding='utf-8') as fi:
            all_errors = fi.readlines()
            all_errors.sort()

        # now write out to final audit file
        with open(details['audit_log'], 'w', encoding='utf-8') as fo:
            for line in all_errors:
                fo.write(f"{line}\n")

        # delete tmp file
        os.remove(details['tmp_audit_log'])

    print('\n\n----------------------------------------------------\n\nProcess complete!')

if __name__ == "__main__":
    
    # check for config file; get the directory where the current script is located
    script_dir = os.path.dirname(os.path.abspath(__file__))

    # Construct the path to the config file; make sure it exists
    config_path = os.path.join(script_dir, 'acf_parse_config.json')
    if os.path.exists(config_path):
        print("\n\nConfig file found.")
    else:
        print(f"\n\nConfig file is missing; please be sure file is saved to {script_dir}.")
        sys.exit(1)

    # load config details
    with open(config_path, 'r') as file:
        details = json.load(file)

    # make sure XSD and XSL files have right paths
    details['xsd_file'] = os.path.join(script_dir, details['xsd_file'])
    details['xsl_file'] = os.path.join(script_dir, details['xsl_file'])

    # make sure files/folders actually exist
    bad_path = False
    for key in ["input_doc", "out_dir", "xsd_file", "xsl_file"]:
        if not os.path.exists(details[key]):
            print(f'\n\n{details[key]} does not exist! Verify path and update config file')
            bad_path = True 
    if bad_path:
        sys.exit(1)

    #create our audit log vars so we can refer to them later
    details['tmp_audit_log'] = os.path.join(details['out_dir'], f'tmp_{details['state'].lower().replace(' ', '_')}_audit-log.txt')
    if os.path.exists(details['tmp_audit_log']):
        os.remove(details['tmp_audit_log'])

    details["audit_log"] = os.path.join(details['out_dir'], f'{details['state'].lower().replace(' ', '_')}_audit-log.txt')
    if os.path.exists(details['audit_log']):
        os.remove(details['audit_log'])

    # make sure boolean values are set
    for term in ["category", "titleContent"]:            
        try:
            if details[term].lower() == 'true':
                details[term] = True 
            elif details[term].lower() == 'false':
                details[term] = False
            else:
                print(f'\n\nThe "{term}" config variable must be "True" or "False".')
                sys.exit(1)
        except KeyError:
            print(f'\n\nThe "{term}" config entry is missing!')
            sys.exit(1)

    #make sure we have regex patterns
    for regex in ["statute_pattern", "state_code_pattern"]:
        if not details[regex]:
            print(f'\n\nMissing the {regex} regex!')
            sys.exit(1)

    # the 'partName' field must be a list, to accommodate variations among records
    if not isinstance(details.get('partName'), list):
        print('\n\nThe "partName" value must be a list!')
        sys.exit(1)

    main(details)

